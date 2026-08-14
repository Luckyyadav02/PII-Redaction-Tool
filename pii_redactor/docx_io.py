"""
docx_io.py
==========
Everything that touches the .docx file format lives here. Detection and
fake-value logic (redactor.py) knows nothing about python-docx, so it
can be unit tested on plain strings.

Known limitation (documented in README): when a paragraph is redacted,
we keep the *first run's* formatting for the whole rewritten paragraph
and blank out the other runs. Mixed-formatting paragraphs (e.g. a name
in bold inside a plain sentence) lose that mid-paragraph styling. Full
run-level diffing was judged not worth the complexity for a redaction
tool -- see README tradeoffs section.
"""

from __future__ import annotations
import csv
from docx import Document

from .redactor import Redactor, RedactionEvent


def _iter_table_paragraphs(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                yield from _iter_table_paragraphs(cell.tables)  # nested tables


def _iter_body_paragraphs(doc):
    yield from doc.paragraphs
    yield from _iter_table_paragraphs(doc.tables)


def _iter_header_footer_paragraphs(doc):
    for section in doc.sections:
        parts = [
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer,
        ]
        for part in parts:
            if part is None:
                continue
            yield from part.paragraphs
            yield from _iter_table_paragraphs(part.tables)


def _redact_paragraph(paragraph, redactor: Redactor) -> list[RedactionEvent]:
    original = paragraph.text
    if not original.strip():
        return []

    new_text, events = redactor.redact_text(original)
    if not events:
        return []

    runs = paragraph.runs
    if runs:
        runs[0].text = new_text
        for extra_run in runs[1:]:
            extra_run.text = ""
    else:
        paragraph.add_run(new_text)

    return events


def redact_docx(input_path: str, output_path: str, redactor: Redactor, log_path: str | None = None):
    """Redact `input_path`, write the result to `output_path`, and
    (optionally) write a CSV audit log of every redaction event to
    `log_path`. Returns (events, counts_by_label)."""
    doc = Document(input_path)

    # Word represents a horizontally/vertically merged cell as separate
    # Cell objects that all wrap the *same* underlying XML paragraph.
    # python-docx's row.cells therefore yields that paragraph once per
    # spanned column/row. Without de-duping by the underlying XML
    # element, a merged cell gets redacted multiple times -- the second
    # pass finds spaCy entities *inside the already-substituted fake
    # text* and redacts those too, corrupting the output.
    seen_elements = set()
    all_events: list[RedactionEvent] = []
    for paragraph in list(_iter_body_paragraphs(doc)) + list(_iter_header_footer_paragraphs(doc)):
        key = id(paragraph._p)
        if key in seen_elements:
            continue
        seen_elements.add(key)
        all_events.extend(_redact_paragraph(paragraph, redactor))

    doc.save(output_path)

    if log_path:
        with open(log_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["label", "original", "fake_replacement"])
            for e in all_events:
                writer.writerow([e.label, e.original, e.fake])

    return all_events, dict(redactor.counts)
